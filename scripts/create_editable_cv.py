from __future__ import annotations

import argparse
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import Color, black
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

from full_stack_cv_content import CV_DATA, CvData, Project, Role


SERIF_FONT = "Times New Roman"
SANS_FONT = "Calibri"
INK = RGBColor(0x11, 0x11, 0x11)
MUTED = RGBColor(0x88, 0x88, 0x88)
CONTENT_WIDTH = Inches(7.4)
PDF_MUTED = Color(0.53, 0.53, 0.53)
WEBSITE_TEXT = CV_DATA.identity.website_text
WEBSITE_URL = CV_DATA.identity.website_url
EMAIL_TEXT = CV_DATA.identity.email
EMAIL_URL = f"mailto:{CV_DATA.identity.email}"
PHONE_TEXT = CV_DATA.identity.phone


@dataclass(frozen=True)
class Entry:
    title_runs: tuple[tuple[str, bool], ...]
    date: str
    descriptions: tuple[str, ...]


PROJECTS = (
    Entry(
        (("Haskell Interpreter", True), ("  ·  Haskell, Yacc", True)),
        "March 2019",
        ("Language and interpreter written in Haskell for processing custom stream data files",),
    ),
    Entry(
        (("GPU Particles", True), ("  ·  Typescript, GLSL", True)),
        "February 2019",
        ("WebGL particle simulation driven by GLSL shaders",),
    ),
    Entry(
        (("Flappy-Pixie", True), ("  ·  Javascript, WebGL", True)),
        "October 2018",
        (
            "Flappy Bird clone with a 3D parallax background. Completed for an interview challenge in one week",
        ),
    ),
    Entry(
        (("Endless-City", True), ("  ·  Javascript, Three.js", True)),
        "September 2018",
        ("Interactive infinite city scene inspired by Little Workshop with custom glTF 2.0 loader",),
    ),
    Entry(
        (("Sushi-Go", True), ("  ·  Java, Swing", True)),
        "April 2018",
        (
            "Multithreaded business software prototype with socket communication and able to save/load configuration files",
        ),
    ),
    Entry(
        (("Fitmed", True), ("  ·  React, Redux, Node.js, MongoDB", True)),
        "July 2018",
        ("Prototype platform for dieticians with API, authentication and input validation",),
    ),
)


def _set_font(run, name: str, size: float, color: RGBColor, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)


def _set_style_font(style, name: str, size: float, color: RGBColor, bold: bool = False) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    r_pr.rFonts.set(qn("w:ascii"), name)
    r_pr.rFonts.set(qn("w:hAnsi"), name)


def add_hyperlink(
    paragraph,
    text: str,
    url: str,
    color: str = "888888",
    size: float = 9,
) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), SANS_FONT)
    fonts.set(qn("w:hAnsi"), SANS_FONT)
    run_properties.append(fonts)

    size_element = OxmlElement("w:sz")
    size_element.set(qn("w:val"), str(int(round(size * 2))))
    run_properties.append(size_element)

    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    run_properties.append(run_color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    run_properties.append(underline)

    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_right_tab(paragraph) -> None:
    paragraph.paragraph_format.tab_stops.add_tab_stop(CONTENT_WIDTH, WD_TAB_ALIGNMENT.RIGHT)


def add_title_date(
    document: Document,
    title_runs: tuple[tuple[str, bool], ...],
    date: str,
    *,
    before: float | None = None,
):
    paragraph = document.add_paragraph(style="CV Entry")
    if before is not None:
        paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.keep_with_next = True
    _add_right_tab(paragraph)

    for text, bold in title_runs:
        run = paragraph.add_run(text)
        _set_font(run, SANS_FONT, 9, INK, bold=bold)

    paragraph.add_run("\t")
    date_run = paragraph.add_run(date)
    _set_font(date_run, SANS_FONT, 9, MUTED)
    return paragraph


def _add_description(
    document: Document,
    text: str,
    *,
    before: float = 0,
    after: float | None = None,
):
    paragraph = document.add_paragraph(style="CV Description")
    paragraph.paragraph_format.space_before = Pt(before)
    if after is not None:
        paragraph.paragraph_format.space_after = Pt(after)
    paragraph.add_run(text)
    return paragraph


def _add_labeled_line(document: Document, label: str, value: str, *, after: float) -> None:
    paragraph = document.add_paragraph(style="CV Skill")
    paragraph.paragraph_format.space_after = Pt(after)
    label_run = paragraph.add_run(label)
    _set_font(label_run, SANS_FONT, 9, INK, bold=True)
    value_run = paragraph.add_run(value)
    _set_font(value_run, SANS_FONT, 9, INK)


def _next_numbering_id(numbering, element_name: str, attribute_name: str) -> int:
    values = []
    for element in numbering.findall(qn(element_name)):
        raw_value = element.get(qn(attribute_name))
        if raw_value is not None:
            values.append(int(raw_value))
    return max(values, default=0) + 1


def _create_bullet_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_id = _next_numbering_id(numbering, "w:abstractNum", "w:abstractNumId")
    number_id = _next_numbering_id(numbering, "w:num", "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    multi_level_type = OxmlElement("w:multiLevelType")
    multi_level_type.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level_type)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "bullet")
    level.append(number_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)

    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "220")
    tabs.append(tab)
    paragraph_properties.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "220")
    indent.set(qn("w:hanging"), "140")
    paragraph_properties.append(indent)
    level.append(paragraph_properties)

    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), SANS_FONT)
    fonts.set(qn("w:hAnsi"), SANS_FONT)
    run_properties.append(fonts)
    level.append(run_properties)
    abstract.append(level)
    numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(number_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    number.append(abstract_reference)
    numbering.append(number)
    return number_id


def _add_bullet(document: Document, text: str, number_id: int) -> None:
    paragraph = document.add_paragraph(style="CV Bullet")
    paragraph_properties = paragraph._p.get_or_add_pPr()
    numbering_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    numbering_properties.append(level)
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(number_id))
    numbering_properties.append(number)
    paragraph_properties.append(numbering_properties)
    paragraph.add_run(text)


def _add_style(document: Document, name: str):
    styles = document.styles
    if name in styles:
        return styles[name]
    return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)

    normal = document.styles["Normal"]
    _set_style_font(normal, SANS_FONT, 9, INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    heading = document.styles["Heading 1"]
    _set_style_font(heading, SERIF_FONT, 16, MUTED)
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(4)
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.keep_together = True

    title = _add_style(document, "CV Title")
    _set_style_font(title, SERIF_FONT, 28, INK)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(0)
    title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    contact = _add_style(document, "CV Contact")
    _set_style_font(contact, SANS_FONT, 9, MUTED)
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(0)
    contact.paragraph_format.line_spacing = Pt(10.5)

    role = _add_style(document, "CV Role")
    _set_style_font(role, SANS_FONT, 11, MUTED, bold=True)
    role.paragraph_format.space_before = Pt(2)
    role.paragraph_format.space_after = Pt(0)
    role.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    entry = _add_style(document, "CV Entry")
    _set_style_font(entry, SANS_FONT, 9, INK, bold=True)
    entry.paragraph_format.space_before = Pt(4)
    entry.paragraph_format.space_after = Pt(0)
    entry.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    entry.paragraph_format.left_indent = Pt(0)

    description = _add_style(document, "CV Description")
    _set_style_font(description, SANS_FONT, 9, INK)
    description.paragraph_format.space_before = Pt(0)
    description.paragraph_format.space_after = Pt(2)
    description.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    description.paragraph_format.left_indent = Pt(0)

    skill = _add_style(document, "CV Skill")
    _set_style_font(skill, SANS_FONT, 9, INK)
    skill.paragraph_format.space_before = Pt(0)
    skill.paragraph_format.space_after = Pt(1.5)
    skill.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    bullet = _add_style(document, "CV Bullet")
    _set_style_font(bullet, SANS_FONT, 9, INK)
    bullet.paragraph_format.space_before = Pt(0)
    bullet.paragraph_format.space_after = Pt(1.2)
    bullet.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    footer = _add_style(document, "CV Footer")
    _set_style_font(footer, SANS_FONT, 8, INK)
    footer.paragraph_format.space_before = Pt(7)
    footer.paragraph_format.space_after = Pt(0)
    footer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _set_cell_margins(cell, value: int = 0) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    margins = cell_properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_properties.append(margins)
    for edge in ("top", "left", "bottom", "right"):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_header_table_geometry(table) -> None:
    column_widths = (6300, 4356)
    table_properties = table._tbl.tblPr

    table_width = table_properties.first_child_found_in("w:tblW")
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(sum(column_widths)))

    table_indent = table_properties.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_properties.append(table_indent)
    table_indent.set(qn("w:type"), "dxa")
    table_indent.set(qn("w:w"), "0")

    grid = table._tbl.tblGrid
    for column in list(grid):
        grid.remove(column)
    for width in column_widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)


def _add_header(document: Document, data: CvData) -> None:
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    left, right = table.rows[0].cells
    left.width = Inches(4.375)
    right.width = Inches(3.025)
    _set_header_table_geometry(table)
    for cell in (left, right):
        _set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    name_paragraph = left.paragraphs[0]
    name_paragraph.style = "CV Title"
    name = name_paragraph.add_run(data.identity.name)
    _set_font(name, SERIF_FONT, 28, INK)

    role_paragraph = left.add_paragraph(style="CV Role")
    role_paragraph.add_run(data.identity.role)

    contact = right.paragraphs[0]
    contact.style = "CV Contact"
    contact.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    contact.paragraph_format.line_spacing = Pt(10.5)
    contact_rows = (
        (
            data.identity.website_text,
            data.identity.website_url,
            True,
        ),
        (
            data.identity.email,
            f"mailto:{data.identity.email}",
            True,
        ),
        (data.identity.phone, "", False),
        (
            data.identity.github_text,
            data.identity.github_url,
            True,
        ),
        (
            data.identity.linkedin_text,
            data.identity.linkedin_url,
            True,
        ),
    )
    for index, (text, url, linked) in enumerate(contact_rows):
        if index:
            line_break = contact.add_run()
            line_break.add_break(WD_BREAK.LINE)
        if linked:
            add_hyperlink(contact, text, url, size=9)
        else:
            phone = contact.add_run(text)
            _set_font(phone, SANS_FONT, 9, MUTED)


def _add_section_heading(document: Document, text: str):
    paragraph = document.add_heading(text, level=1)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def _add_role(document: Document, role: Role, bullet_number_id: int) -> None:
    add_title_date(document, ((role.title, True),), role.period)
    for bullet in role.bullets:
        _add_bullet(document, bullet, bullet_number_id)


def _add_project(document: Document, project: Project) -> None:
    add_title_date(
        document,
        ((project.title, True), (f"  ·  {project.tools}", True)),
        project.period,
    )
    paragraph = _add_description(document, project.description)
    if project.url:
        paragraph.add_run("  ")
        add_hyperlink(
            paragraph,
            "View project",
            project.url,
            color="555555",
            size=8.5,
        )


def build_cv(output_path: Path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    configure_document(document)
    document.core_properties.title = "Tomasz Zielinski - CV"
    document.core_properties.author = "Tomasz Zielinski"
    document.core_properties.subject = "Curriculum Vitae"

    bullet_number_id = _create_bullet_numbering(document)
    _add_header(document, CV_DATA)

    _add_section_heading(document, "Profile")
    _add_description(document, CV_DATA.profile)

    _add_section_heading(document, "Core Technologies")
    for label, technologies in CV_DATA.technology_groups:
        _add_labeled_line(
            document,
            f"{label}:  ",
            ", ".join(technologies),
            after=1.5,
        )

    _add_section_heading(document, "Professional Experience")
    add_title_date(
        document,
        ((CV_DATA.employment.company, True),),
        CV_DATA.employment.period,
    )
    for role in CV_DATA.employment.roles:
        _add_role(document, role, bullet_number_id)

    _add_section_heading(document, "Selected Product Work")
    for project in CV_DATA.commercial_projects:
        _add_project(document, project)

    document.add_page_break()

    _add_section_heading(document, "Earlier Experience")
    _add_description(document, CV_DATA.earlier_experience[0])
    _add_description(document, CV_DATA.earlier_experience[1])

    _add_section_heading(document, "Selected Projects")
    for project in CV_DATA.personal_projects:
        _add_project(document, project)

    _add_section_heading(document, "Education")
    for education in CV_DATA.education:
        add_title_date(
            document,
            ((education.institution, True),),
            education.period,
        )
        _add_description(document, education.qualification)

    _add_section_heading(document, "Additional Information")
    _add_labeled_line(
        document,
        "Languages:  ",
        ", ".join(CV_DATA.languages),
        after=1.5,
    )
    consent = document.add_paragraph(style="CV Footer")
    consent.add_run(CV_DATA.consent)

    document.save(output_path)
    return output_path


PDF_SERIF = "CvTimes"
PDF_SANS = "CvCalibri"
PDF_SANS_BOLD = "CvCalibriBold"


def _register_pdf_fonts() -> None:
    registrations = (
        (PDF_SERIF, Path(r"C:\Windows\Fonts\times.ttf")),
        (PDF_SANS, Path(r"C:\Windows\Fonts\calibri.ttf")),
        (PDF_SANS_BOLD, Path(r"C:\Windows\Fonts\calibrib.ttf")),
    )
    registered = set(pdfmetrics.getRegisteredFontNames())
    for name, path in registrations:
        if name not in registered:
            pdfmetrics.registerFont(TTFont(name, str(path)))


def _pdf_header(pdf: canvas.Canvas, data: CvData) -> float:
    pdf.setFillColor(black)
    pdf.setFont(PDF_SERIF, 28)
    pdf.drawString(40, 744, data.identity.name)
    pdf.setFont(PDF_SANS_BOLD, 11)
    pdf.drawString(40, 724, data.identity.role)

    contact_rows = (
        (data.identity.website_text, data.identity.website_url, 748),
        (data.identity.email, f"mailto:{data.identity.email}", 735),
        (data.identity.phone, "", 722),
        (data.identity.github_text, data.identity.github_url, 709),
        (data.identity.linkedin_text, data.identity.linkedin_url, 696),
    )
    pdf.setFont(PDF_SANS, 9)
    pdf.setFillColor(PDF_MUTED)
    for text, url, y in contact_rows:
        pdf.drawRightString(572, y, text)
        if url:
            width = stringWidth(text, PDF_SANS, 9)
            pdf.linkURL(url, (572 - width, y - 1, 572, y + 9), relative=0)
    return 674


def _pdf_section(pdf: canvas.Canvas, title: str, y: float) -> float:
    pdf.setFillColor(PDF_MUTED)
    pdf.setFont(PDF_SERIF, 16)
    pdf.drawString(40, y, title)
    return y - 25


def _pdf_wrapped_text(
    pdf: canvas.Canvas,
    text: str,
    x: float,
    y: float,
    width: float,
    font_name: str = PDF_SANS,
    size: float = 9,
    leading: float = 11,
) -> float:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if current and stringWidth(candidate, font_name, size) > width:
            lines.append(current)
            current = word
        else:
            current = candidate
    if current:
        lines.append(current)

    pdf.setFont(font_name, size)
    pdf.setFillColor(black)
    for line in lines:
        pdf.drawString(x, y, line)
        y -= leading
    return y


def _pdf_title_date(
    pdf: canvas.Canvas,
    title: str,
    date: str,
    y: float,
    *,
    x: float = 40,
) -> float:
    pdf.setFillColor(black)
    pdf.setFont(PDF_SANS_BOLD, 9)
    pdf.drawString(x, y, title)
    pdf.setFillColor(PDF_MUTED)
    pdf.setFont(PDF_SANS, 9)
    pdf.drawRightString(572, y, date)
    return y - 12


def _pdf_labeled_line(
    pdf: canvas.Canvas,
    label: str,
    value: str,
    y: float,
) -> float:
    pdf.setFillColor(black)
    pdf.setFont(PDF_SANS_BOLD, 9)
    pdf.drawString(40, y, label)
    label_width = stringWidth(label, PDF_SANS_BOLD, 9)
    return _pdf_wrapped_text(
        pdf,
        value,
        x=40 + label_width,
        y=y,
        width=532 - label_width,
        font_name=PDF_SANS,
        size=9,
        leading=11,
    )


def _pdf_bullet(pdf: canvas.Canvas, text: str, y: float) -> float:
    pdf.setFillColor(black)
    pdf.setFont(PDF_SANS, 9)
    pdf.drawString(41, y, "•")
    next_y = _pdf_wrapped_text(
        pdf,
        text,
        x=49,
        y=y,
        width=519,
        font_name=PDF_SANS,
        size=9,
        leading=11,
    )
    return next_y - 2


def _pdf_project(pdf: canvas.Canvas, project: Project, y: float) -> float:
    y = _pdf_title_date(pdf, project.title, project.period, y)
    y = _pdf_wrapped_text(
        pdf,
        project.tools,
        x=40,
        y=y,
        width=532,
        font_name=PDF_SANS_BOLD,
        size=8.5,
        leading=10,
    )
    y = _pdf_wrapped_text(pdf, project.description, 40, y, 532)
    if project.url:
        pdf.setFillColor(PDF_MUTED)
        pdf.setFont(PDF_SANS, 8)
        label = "View project"
        pdf.drawString(40, y, label)
        width = stringWidth(label, PDF_SANS, 8)
        pdf.linkURL(project.url, (40, y - 1, 40 + width, y + 8), relative=0)
        y -= 10
    return y - 4


def build_pdf(output_path: Path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    _register_pdf_fonts()

    pdf = canvas.Canvas(str(output_path), pagesize=letter, pageCompression=1)
    pdf.setTitle("Tomasz Zielinski - Full-Stack Engineer CV")
    pdf.setAuthor(CV_DATA.identity.name)
    pdf.setSubject("Curriculum Vitae")

    y = _pdf_header(pdf, CV_DATA)
    y = _pdf_section(pdf, "Profile", y)
    y = _pdf_wrapped_text(pdf, CV_DATA.profile, 40, y, 532) - 7

    y = _pdf_section(pdf, "Core Technologies", y)
    for label, technologies in CV_DATA.technology_groups:
        y = _pdf_labeled_line(pdf, f"{label}:  ", ", ".join(technologies), y) - 2
    y -= 4

    y = _pdf_section(pdf, "Professional Experience", y)
    y = _pdf_title_date(
        pdf,
        CV_DATA.employment.company,
        CV_DATA.employment.period,
        y,
    )
    for role in CV_DATA.employment.roles:
        y = _pdf_title_date(pdf, role.title, role.period, y)
        for bullet in role.bullets:
            y = _pdf_bullet(pdf, bullet, y)
        y -= 2

    y = _pdf_section(pdf, "Selected Product Work", y)
    for project in CV_DATA.commercial_projects:
        y = _pdf_project(pdf, project, y)

    pdf.showPage()

    y = 744
    y = _pdf_section(pdf, "Earlier Experience", y)
    y = _pdf_wrapped_text(
        pdf,
        CV_DATA.earlier_experience[0],
        40,
        y,
        532,
        font_name=PDF_SANS_BOLD,
    )
    y = _pdf_wrapped_text(pdf, CV_DATA.earlier_experience[1], 40, y, 532) - 5

    y = _pdf_section(pdf, "Selected Projects", y)
    for project in CV_DATA.personal_projects:
        y = _pdf_project(pdf, project, y)

    y = _pdf_section(pdf, "Education", y)
    for education in CV_DATA.education:
        y = _pdf_title_date(pdf, education.institution, education.period, y)
        y = _pdf_wrapped_text(pdf, education.qualification, 40, y, 532) - 3

    y = _pdf_section(pdf, "Additional Information", y)
    y = _pdf_labeled_line(pdf, "Languages:  ", ", ".join(CV_DATA.languages), y) - 6
    _pdf_wrapped_text(
        pdf,
        CV_DATA.consent,
        40,
        y,
        532,
        font_name=PDF_SANS,
        size=8,
        leading=9.5,
    )

    pdf.showPage()
    pdf.save()
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Create the editable Tomasz Zielinski CV.")
    parser.add_argument("output", type=Path, help="Destination DOCX path")
    parser.add_argument("--pdf", action="store_true", help="Create a PDF instead of a DOCX")
    arguments = parser.parse_args()
    if arguments.pdf:
        build_pdf(arguments.output)
    else:
        build_cv(arguments.output)


if __name__ == "__main__":
    main()
