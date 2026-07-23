from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path
from xml.etree import ElementTree
from zipfile import ZipFile

from docx import Document
from docx.shared import Inches, Pt
import pdfplumber
from pypdf import PdfReader

sys.path.insert(0, str(Path(__file__).resolve().parent))

from create_editable_cv import build_cv, build_pdf
from full_stack_cv_content import CV_DATA, FORBIDDEN_PUBLIC_TERMS


class EditableCvBuilderTest(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls._temp_dir = tempfile.TemporaryDirectory()
        cls.output_path = Path(cls._temp_dir.name) / "cv.docx"
        build_cv(cls.output_path)
        cls.document = Document(cls.output_path)
        body_text = [paragraph.text for paragraph in cls.document.paragraphs]
        table_text = [
            paragraph.text
            for table in cls.document.tables
            for row in table.rows
            for cell in row.cells
            for paragraph in cell.paragraphs
        ]
        cls.text = "\n".join(table_text + body_text)

    @classmethod
    def tearDownClass(cls) -> None:
        cls._temp_dir.cleanup()

    def test_profile_contains_the_approved_ownership_sentence(self) -> None:
        self.assertIn(
            "I translate operational requirements into maintainable systems and "
            "take ownership across the delivery lifecycle.",
            CV_DATA.profile,
        )

    def test_builds_letter_document_with_compact_one_page_margins(self) -> None:
        section = self.document.sections[0]

        self.assertEqual(section.page_width, Inches(8.5))
        self.assertEqual(section.page_height, Inches(11))
        self.assertAlmostEqual(section.top_margin, Inches(0.42), delta=Pt(0.1))
        self.assertAlmostEqual(section.bottom_margin, Inches(0.4), delta=Pt(0.1))
        self.assertAlmostEqual(section.left_margin, Inches(0.5), delta=Pt(0.1))
        self.assertAlmostEqual(section.right_margin, Inches(0.5), delta=Pt(0.1))

    def test_preserves_section_order_and_visible_content(self) -> None:
        headings = [
            paragraph.text
            for paragraph in self.document.paragraphs
            if paragraph.style.name == "Heading 1"
        ]

        self.assertEqual(
            headings,
            [
                "Profile",
                "Education",
                "Professional Experience",
                "Commercial Experience",
                "Technologies",
                "Selected Projects",
                "Additional Information",
            ],
        )

        expected_fragments = [
            "Tomasz Zielinski",
            "Full-Stack Developer",
            "corashina.github.io",
            "corashina@gmail.com",
            "+48 791 748 226",
            "Profile",
            "Education",
            "University of Southampton",
            "July 2020",
            "Xelto",
            "Freelance Web Development",
            "Commercial Experience",
            "Technologies",
            "Selected Projects",
            "Endless City",
            "Flappy-Pixie",
            "Fitmed",
            "Additional Information",
            "I hereby consent to the processing of personal data",
        ]

        positions = [self.text.index(fragment) for fragment in expected_fragments]
        self.assertEqual(positions, sorted(positions))

        for obsolete_contact in (
            "www.zielin.ski",
            "contact@zielin.ski",
            "07519554924",
        ):
            self.assertNotIn(obsolete_contact, self.text)

    def test_uses_named_styles_and_editable_word_structures(self) -> None:
        required_styles = {
            "CV Title",
            "CV Role",
            "CV Company",
            "CV Entry",
            "CV Description",
            "CV Bullet",
            "CV Footer",
        }
        style_names = {style.name for style in self.document.styles}
        self.assertTrue(required_styles.issubset(style_names))

        with ZipFile(self.output_path) as archive:
            document_xml = archive.read("word/document.xml")
            relationships_xml = archive.read("word/_rels/document.xml.rels")

        self.assertIn(b'w:val="right"', document_xml)
        self.assertIn(b"<w:numPr>", document_xml)
        self.assertIn(b"<w:hyperlink", document_xml)
        self.assertEqual(len(self.document.tables), 1)
        self.assertEqual(len(self.document.tables[0].rows), 1)
        self.assertEqual(len(self.document.tables[0].columns), 2)
        self.assertNotIn(b"<w:txbxContent", document_xml)
        self.assertIn(b"relationships/hyperlink", relationships_xml)

    def test_uses_compact_one_page_font_roles(self) -> None:
        expected_styles = {
            "CV Title": ("Times New Roman", 26.0),
            "Heading 1": ("Times New Roman", 13.5),
            "CV Contact": ("Calibri", 8.5),
            "CV Role": ("Calibri", 11.0),
            "CV Company": ("Times New Roman", 10.5),
            "CV Entry": ("Calibri", 9.0),
            "CV Description": ("Calibri", 8.5),
            "CV Bullet": ("Calibri", 8.5),
            "CV Skill": ("Calibri", 8.5),
            "CV Footer": ("Calibri", 8.5),
        }

        for style_name, (font_name, font_size) in expected_styles.items():
            style = self.document.styles[style_name]
            self.assertEqual(style.font.name, font_name)
            self.assertAlmostEqual(style.font.size.pt, font_size, delta=0.01)

    def test_uses_compact_body_fonts_and_keep_rules(self) -> None:
        for style_name in ("CV Description", "CV Bullet", "CV Skill", "CV Footer"):
            style = self.document.styles[style_name]
            self.assertAlmostEqual(style.font.size.pt, 8.5, delta=0.01)
        heading_style = self.document.styles["Heading 1"]
        self.assertAlmostEqual(
            heading_style.paragraph_format.space_before.pt,
            8.5,
            delta=0.01,
        )
        for paragraph in self.document.paragraphs:
            if paragraph.style.name == "Heading 1":
                self.assertTrue(paragraph.paragraph_format.keep_with_next)

    def test_has_no_placeholders_or_extraction_artifacts(self) -> None:
        for forbidden in ("TBD", "TODO", "\u200b", "\ufffd"):
            self.assertNotIn(forbidden, self.text)

    def test_full_stack_content_model_contains_approved_current_facts(self) -> None:
        self.assertEqual(CV_DATA.identity.role, "Full-Stack Developer")
        self.assertEqual(CV_DATA.identity.phone, "+48 791 748 226")
        self.assertEqual(CV_DATA.identity.email, "corashina@gmail.com")
        self.assertEqual(CV_DATA.identity.website_text, "corashina.github.io")
        self.assertEqual(
            [(role.title, role.period) for role in CV_DATA.employment.roles],
            [
                ("Full-Stack Developer", "2024–2026"),
                ("Junior Frontend Developer", "2021–2024"),
            ],
        )
        self.assertEqual(
            [project.title for project in CV_DATA.personal_projects],
            ["Endless City", "Flappy-Pixie", "Fitmed"],
        )
        self.assertEqual(CV_DATA.education[0].period, "July 2020")

    def test_public_content_excludes_private_metrics_and_client_names(self) -> None:
        visible = CV_DATA.visible_text()
        for forbidden in FORBIDDEN_PUBLIC_TERMS:
            self.assertNotIn(forbidden.casefold(), visible.casefold())

    def test_docx_uses_one_page_structure_without_social_links(self) -> None:
        with ZipFile(self.output_path) as archive:
            document_xml = archive.read("word/document.xml")
            relationships_xml = archive.read("word/_rels/document.xml.rels")

        self.assertEqual(document_xml.count(b'w:type="page"'), 0)
        self.assertNotIn(b"github.com/corashina", document_xml)
        self.assertNotIn(b"linkedin.com/in/", document_xml)
        self.assertNotIn(b"View project", document_xml)
        self.assertNotIn(b"github.com/corashina", relationships_xml)
        relationships = ElementTree.fromstring(relationships_xml)
        external_targets = {
            relationship.attrib["Target"]
            for relationship in relationships
            if relationship.attrib.get("TargetMode") == "External"
        }
        self.assertEqual(
            external_targets,
            {
                CV_DATA.identity.website_url,
                f"mailto:{CV_DATA.identity.email}",
            },
        )

    def test_docx_distinguishes_company_from_role(self) -> None:
        company_paragraph = next(
            paragraph for paragraph in self.document.paragraphs
            if paragraph.text.startswith("Xelto")
        )
        role_paragraph = next(
            paragraph for paragraph in self.document.paragraphs
            if paragraph.text.startswith("Full-Stack Developer")
        )
        self.assertEqual(company_paragraph.style.name, "CV Company")
        self.assertEqual(role_paragraph.style.name, "CV Entry")
        self.assertNotEqual(
            company_paragraph.style.font.name,
            role_paragraph.style.font.name,
        )
        self.assertGreater(
            company_paragraph.style.font.size.pt,
            role_paragraph.style.font.size.pt,
        )

    def test_freelance_work_is_inside_professional_experience(self) -> None:
        texts = [paragraph.text for paragraph in self.document.paragraphs]
        professional = texts.index("Professional Experience")
        commercial = texts.index("Commercial Experience")
        freelance = next(
            index for index, text in enumerate(texts)
            if text.startswith("Freelance Web Development")
        )
        self.assertLess(professional, freelance)
        self.assertLess(freelance, commercial)
        self.assertNotIn("Earlier Experience", texts)

    def test_builds_one_page_full_stack_docx_structure(self) -> None:
        headings = [
            paragraph.text
            for paragraph in self.document.paragraphs
            if paragraph.style.name == "Heading 1"
        ]
        self.assertEqual(
            headings,
            [
                "Profile",
                "Education",
                "Professional Experience",
                "Commercial Experience",
                "Technologies",
                "Selected Projects",
                "Additional Information",
            ],
        )
        with ZipFile(self.output_path) as archive:
            document_xml = archive.read("word/document.xml")
        self.assertEqual(document_xml.count(b'w:type="page"'), 0)
        self.assertIn(b"Full-Stack Developer", document_xml)
        self.assertNotIn(b"Full-Stack Engineer", document_xml)
        self.assertNotIn(b"expected in July 2020", document_xml)

    def test_docx_has_only_approved_contact_rows_and_projects(self) -> None:
        contact_rows = tuple(
            self.document.tables[0].rows[0].cells[1].paragraphs[0].text.splitlines()
        )
        self.assertEqual(
            contact_rows,
            ("corashina.github.io", "corashina@gmail.com", "+48 791 748 226"),
        )
        for value in (
            "corashina.github.io",
            "corashina@gmail.com",
            "+48 791 748 226",
            "Endless City",
            "Flappy-Pixie",
            "Fitmed",
        ):
            self.assertIn(value, self.text)
        for obsolete in (
            "github.com/corashina",
            "linkedin.com/in/tomasz-zielinski-a97999161",
            "Haskell Interpreter",
            "GPU Particles",
            "Sushi-Go",
        ):
            self.assertNotIn(obsolete, self.text)

    def test_builds_one_page_pdf_with_approved_hierarchy_and_links(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            output = Path(temporary_directory) / "cv.pdf"
            build_pdf(output)
            reader = PdfReader(output)
            self.assertEqual(len(reader.pages), 1)
            page = reader.pages[0]
            self.assertEqual(float(page.mediabox.width), 612)
            self.assertEqual(float(page.mediabox.height), 792)
            text = page.extract_text() or ""
            for heading in (
                "Profile",
                "Education",
                "Professional Experience",
                "Commercial Experience",
                "Technologies",
                "Selected Projects",
                "Additional Information",
            ):
                self.assertIn(heading, text)
            for removed in (
                "Core Technologies",
                "Selected Product Work",
                "Earlier Experience",
                "github.com/corashina",
                "linkedin.com/in/",
                "View project",
            ):
                self.assertNotIn(removed, text)

            annotations = page.get("/Annots", [])
            uris = [
                annotation.get_object().get("/A", {}).get("/URI", "")
                for annotation in annotations
            ]
            self.assertEqual(
                set(uris),
                {
                    CV_DATA.identity.website_url,
                    f"mailto:{CV_DATA.identity.email}",
                },
            )
            for removed_uri in (
                CV_DATA.identity.github_url,
                CV_DATA.identity.linkedin_url,
                *(project.url for project in CV_DATA.personal_projects),
            ):
                self.assertNotIn(removed_uri, uris)
            for forbidden in FORBIDDEN_PUBLIC_TERMS:
                self.assertNotIn(forbidden.casefold(), text.casefold())

            with pdfplumber.open(output) as pdf:
                rendered_page = pdf.pages[0]
                words = rendered_page.extract_words()
            self.assertTrue(words)
            self.assertLessEqual(
                max(word["bottom"] for word in words),
                rendered_page.height - 28.8,
            )
            self.assertGreaterEqual(
                max(word["bottom"] for word in words),
                rendered_page.height - 34,
            )


if __name__ == "__main__":
    unittest.main(verbosity=2)
